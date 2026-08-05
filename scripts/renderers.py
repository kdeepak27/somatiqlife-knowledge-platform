from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor

from schemas.module_schema import ModuleContent


def safe_filename(value: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*]', "", value)
    return re.sub(r"\s+", "_", cleaned.strip())


def render_json(module: ModuleContent, destination: Path) -> Path:
    destination.write_text(
        module.model_dump_json(indent=2),
        encoding="utf-8",
    )
    return destination


def render_markdown(module: ModuleContent, destination: Path) -> Path:
    lines: list[str] = [
        f"# {module.module_id} - {module.title}",
        f"## {module.subtitle}",
        "",
        f"**Domain:** {module.domain}",
        f"**Category:** {module.category}",
        f"**SOM stages:** {', '.join(module.som_stages)}",
        f"**Audience:** {', '.join(module.audience)}",
        f"**Reading level:** {module.reading_level}",
        f"**Keywords:** {', '.join(module.keywords)}",
        "",
        "## AI Retrieval Summary",
        module.retrieval_summary,
        "",
        "## Customer Questions Answered",
    ]

    lines.extend(f"- {item}" for item in module.customer_questions)
    lines.extend(["", "## Learning Objectives"])
    lines.extend(f"- {item}" for item in module.learning_objectives)

    lines.extend(
        [
            "",
            "## Definition",
            module.definition,
            "",
            "## Why This Matters",
            module.why_this_matters,
            "",
        ]
    )

    for section in module.sections:
        lines.extend([f"## {section.heading}", section.content, ""])

    lines.append("## Practical Guidance")
    lines.extend(f"- {item}" for item in module.practical_guidance)

    lines.extend(["", "## Safety Notes"])
    lines.extend(f"- {item}" for item in module.safety_notes)

    lines.extend(["", "## Common Misconceptions"])
    for item in module.misconceptions:
        lines.extend(
            [
                f"### Misconception: {item.misconception}",
                f"**Reality:** {item.reality}",
                "",
            ]
        )

    lines.append("## Clinical Pearls")
    lines.extend(f"- {item}" for item in module.clinical_pearls)

    lines.extend(["", "## Key Takeaways"])
    lines.extend(f"- {item}" for item in module.key_takeaways)

    lines.extend(["", "## Glossary"])
    for term, definition in module.glossary.items():
        lines.append(f"- **{term}:** {definition}")

    lines.extend(["", "## Related Modules"])
    lines.extend(f"- {item}" for item in module.related_modules)

    lines.extend(["", "## Illustration Specifications"])
    for figure in module.illustrations:
        lines.extend(
            [
                f"### {figure.illustration_id} - {figure.title}",
                f"**Type:** {figure.illustration_type}",
                "",
                figure.description,
                "",
                f"**Alt text:** {figure.alt_text}",
                "",
            ]
        )

    lines.append("## References")
    for reference in module.references:
        fields = [
            reference.title,
            reference.organization_or_authors,
            reference.year,
            reference.identifier,
            reference.verification_status,
        ]
        lines.append("- " + " | ".join(field for field in fields if field))

    lines.extend(
        [
            "",
            f"**Evidence strength:** {module.evidence_strength}",
            f"**Version:** {module.version}",
            f"**Review status:** {module.review_status}",
        ]
    )

    destination.write_text("\n".join(lines), encoding="utf-8")
    return destination


def render_docx(module: ModuleContent, destination: Path) -> Path:
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    document.styles["Normal"].font.name = "Aptos"
    document.styles["Normal"].font.size = Pt(10.5)

    for style_name, size, color in [
        ("Title", 26, RGBColor(183, 65, 14)),
        ("Heading 1", 17, RGBColor(183, 65, 14)),
        ("Heading 2", 13, RGBColor(92, 45, 28)),
    ]:
        style = document.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = color

    document.add_heading(f"{module.module_id} - {module.title}", level=0)

    subtitle = document.add_paragraph()
    subtitle.add_run(module.subtitle).italic = True

    metadata = document.add_table(rows=0, cols=2)
    metadata.style = "Table Grid"
    rows = [
        ("Domain", module.domain),
        ("Category", module.category),
        ("SOM stages", ", ".join(module.som_stages)),
        ("Audience", ", ".join(module.audience)),
        ("Reading level", module.reading_level),
        ("Keywords", ", ".join(module.keywords)),
        ("Version", module.version),
        ("Review status", module.review_status),
    ]
    for key, value in rows:
        cells = metadata.add_row().cells
        cells[0].text = key
        cells[1].text = value
        cells[0].paragraphs[0].runs[0].bold = True

    document.add_heading("AI Retrieval Summary", level=1)
    document.add_paragraph(module.retrieval_summary)

    document.add_heading("Customer Questions Answered", level=1)
    for item in module.customer_questions:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Learning Objectives", level=1)
    for item in module.learning_objectives:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Definition", level=1)
    document.add_paragraph(module.definition)

    document.add_heading("Why This Matters", level=1)
    document.add_paragraph(module.why_this_matters)

    for item in module.sections:
        document.add_heading(item.heading, level=1)
        document.add_paragraph(item.content)

    document.add_heading("Practical Guidance", level=1)
    for item in module.practical_guidance:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Safety Notes", level=1)
    for item in module.safety_notes:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Common Misconceptions", level=1)
    for item in module.misconceptions:
        p = document.add_paragraph()
        p.add_run("Misconception: ").bold = True
        p.add_run(item.misconception)
        p = document.add_paragraph()
        p.add_run("Reality: ").bold = True
        p.add_run(item.reality)

    document.add_heading("Clinical Pearls", level=1)
    for item in module.clinical_pearls:
        document.add_paragraph(item, style="List Number")

    document.add_heading("Key Takeaways", level=1)
    for item in module.key_takeaways:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Glossary", level=1)
    for term, definition in module.glossary.items():
        p = document.add_paragraph()
        p.add_run(f"{term}: ").bold = True
        p.add_run(definition)

    document.add_heading("Related Modules", level=1)
    for item in module.related_modules:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Illustration Specifications", level=1)
    for figure in module.illustrations:
        document.add_heading(
            f"{figure.illustration_id} - {figure.title}",
            level=2,
        )
        document.add_paragraph(f"Type: {figure.illustration_type}")
        document.add_paragraph(figure.description)
        document.add_paragraph(f"Alt text: {figure.alt_text}")

    document.add_heading("References", level=1)
    for reference in module.references:
        fields = [
            reference.title,
            reference.organization_or_authors,
            reference.year,
            reference.identifier,
            reference.verification_status,
        ]
        document.add_paragraph(
            " | ".join(field for field in fields if field),
            style="List Bullet",
        )

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = 1
    run = footer.add_run(
        "SOMatiq Knowledge Platform | Draft | Review required"
    )
    run.font.size = Pt(8)

    document.save(destination)
    return destination


def render_jsonl_chunks(module: ModuleContent, destination: Path) -> Path:
    base = {
        "module_id": module.module_id,
        "module_title": module.title,
        "domain": module.domain,
        "category": module.category,
        "keywords": module.keywords,
        "version": module.version,
        "review_status": module.review_status,
    }

    records = [
        {
            **base,
            "chunk_id": f"{module.module_id}-SUMMARY",
            "chunk_title": "AI Retrieval Summary",
            "content": module.retrieval_summary,
        },
        {
            **base,
            "chunk_id": f"{module.module_id}-DEFINITION",
            "chunk_title": "Definition",
            "content": module.definition,
        },
        {
            **base,
            "chunk_id": f"{module.module_id}-WHY-IT-MATTERS",
            "chunk_title": "Why This Matters",
            "content": module.why_this_matters,
        },
    ]

    for index, section in enumerate(module.sections, start=1):
        records.append(
            {
                **base,
                "chunk_id": f"{module.module_id}-SECTION-{index:02d}",
                "chunk_title": section.heading,
                "content": section.content,
            }
        )

    with destination.open("w", encoding="utf-8") as handle:
        for record in records:
            handle.write(json.dumps(record, ensure_ascii=False) + "\n")

    return destination
